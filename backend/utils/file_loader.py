"""
=============================================================================
HACK HUNTERS - File Loader Utility (Production-Ready)
=============================================================================
Robust text extraction from PDF, DOCX, and TXT files.

PDF extraction cascade:
    1. PyPDF2 (fast, native text)
    2. pdfplumber (better layout handling)
    3. OCR via pdf2image + pytesseract (scanned / image-based PDFs)

Handles:
    - Corrupted files
    - Password-protected PDFs
    - Scanned / image-based PDFs
    - Empty documents
    - Encoding issues (TXT)
=============================================================================
"""

import logging
from pathlib import Path
from typing import Optional

logger = logging.getLogger(__name__)


# ── Minimum characters to consider extraction "successful" ──────────────────
_MIN_TEXT_LENGTH = 50  # Reduced from 100 to handle shorter documents


# ═══════════════════════════════════════════════════════════════════════════════
#  Custom Exception Hierarchy
# ═══════════════════════════════════════════════════════════════════════════════

class FileLoadError(Exception):
    """Base class for file loading errors.  Always carries a user-facing message."""
    pass


class UnsupportedFileTypeError(FileLoadError):
    """Raised when the file extension is not in the allowed set."""
    pass


class CorruptedFileError(FileLoadError):
    """Raised when the file appears to be corrupted or unreadable."""
    pass


class PasswordProtectedError(FileLoadError):
    """Raised when a PDF is password-protected / encrypted."""
    pass


class EmptyDocumentError(FileLoadError):
    """Raised when the document contains no extractable text."""
    pass


class OCRUnavailableError(FileLoadError):
    """Raised when OCR dependencies are not installed."""
    pass


# ═══════════════════════════════════════════════════════════════════════════════
#  FileLoader
# ═══════════════════════════════════════════════════════════════════════════════

class FileLoader:
    """
    Extracts text content from supported document formats.

    Supported formats:
        - PDF  (.pdf)  via PyPDF2 -> pdfplumber -> OCR
        - DOCX (.docx) via python-docx
        - TXT  (.txt)  via standard file I/O
    """

    SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".txt"}

    # ── Public API ──────────────────────────────────────────────────────────

    @staticmethod
    def validate_extension(file_path: str) -> str:
        """
        Validate file extension and return normalised extension.

        Args:
            file_path: Path or file name.

        Returns:
            Lowercase file extension (e.g., '.pdf').

        Raises:
            UnsupportedFileTypeError: If the extension is not supported.
        """
        ext = Path(file_path).suffix.lower()
        if ext not in FileLoader.SUPPORTED_EXTENSIONS:
            raise UnsupportedFileTypeError(
                f"Unsupported file type: '{ext}'. "
                f"Supported types: {', '.join(sorted(FileLoader.SUPPORTED_EXTENSIONS))}"
            )
        return ext

    @staticmethod
    def load(file_path: str) -> str:
        """
        Load and extract text from a document.

        Args:
            file_path: Absolute or relative path to the document file.

        Returns:
            Extracted text content as a string (guaranteed non-empty).

        Raises:
            FileLoadError (or subclass): On any failure.
        """
        path = Path(file_path)

        if not path.exists():
            raise FileLoadError(f"File not found: {file_path}")

        if not path.is_file():
            raise FileLoadError(f"Path is not a file: {file_path}")

        ext = FileLoader.validate_extension(file_path)

        try:
            if ext == ".pdf":
                return FileLoader._load_pdf(path)
            elif ext == ".docx":
                return FileLoader._load_docx(path)
            elif ext == ".txt":
                return FileLoader._load_txt(path)
            else:
                raise UnsupportedFileTypeError(f"No loader for extension: {ext}")
        except FileLoadError:
            raise
        except Exception as e:
            logger.error(f"Unexpected error loading file '{file_path}': {e}", exc_info=True)
            raise CorruptedFileError(
                f"The uploaded file appears to be corrupted or unreadable."
            )

    # ── PDF Extraction (3-stage cascade) ────────────────────────────────────

    @staticmethod
    def _load_pdf(path: Path) -> str:
        """
        Extract text from a PDF using a 3-stage cascade:
            1. PyPDF2   (fast, native digital text)
            2. pdfplumber (better layout/table handling)
            3. OCR      (scanned / image-based pages)

        Raises specific exceptions for password-protected, corrupted,
        and empty PDFs.
        """
        logger.info(f"PDF extraction starting: {path.name}")

        # ── Stage 0: Detect password-protection / corruption early ──────
        FileLoader._pdf_preflight_check(path)

        # ── Stage 1: PyPDF2 ─────────────────────────────────────────────
        text = FileLoader._pdf_stage_pypdf2(path)
        if text and len(text.strip()) >= _MIN_TEXT_LENGTH:
            logger.info(
                f"PDF '{path.name}': PyPDF2 extracted {len(text)} chars (stage 1)."
            )
            return text.strip()

        logger.info(
            f"PDF '{path.name}': PyPDF2 yielded {len(text.strip()) if text else 0} chars "
            f"(< {_MIN_TEXT_LENGTH}). Trying pdfplumber..."
        )

        # ── Stage 2: pdfplumber ─────────────────────────────────────────
        text = FileLoader._pdf_stage_pdfplumber(path)
        if text and len(text.strip()) >= _MIN_TEXT_LENGTH:
            logger.info(
                f"PDF '{path.name}': pdfplumber extracted {len(text)} chars (stage 2)."
            )
            return text.strip()

        logger.info(
            f"PDF '{path.name}': pdfplumber yielded {len(text.strip()) if text else 0} chars "
            f"(< {_MIN_TEXT_LENGTH}). Trying OCR..."
        )

        # ── Stage 3: OCR (pdf2image + pytesseract) ─────────────────────
        text = FileLoader._pdf_stage_ocr(path)
        if text and len(text.strip()) >= _MIN_TEXT_LENGTH:
            logger.info(
                f"PDF '{path.name}': OCR extracted {len(text)} chars (stage 3)."
            )
            return text.strip()

        # ── All stages failed ───────────────────────────────────────────
        raise EmptyDocumentError(
            "This PDF appears to contain no readable text. "
            "It may be a scanned document that OCR could not process, "
            "or the pages may contain only images/graphics."
        )

    @staticmethod
    def _pdf_preflight_check(path: Path) -> None:
        """
        Quick pre-flight: detect encrypted / corrupted PDFs before
        running the full extraction pipeline.
        """
        try:
            from PyPDF2 import PdfReader
            from PyPDF2.errors import PdfReadError

            reader = PdfReader(str(path))

            # PyPDF2 sets is_encrypted = True for password-protected files
            if reader.is_encrypted:
                # Try empty password (some PDFs are "encrypted" with blank password)
                try:
                    decrypted = reader.decrypt("")
                    if decrypted == 0:
                        raise PasswordProtectedError(
                            "Password-protected PDFs are not supported. "
                            "Please remove the password and re-upload."
                        )
                except Exception:
                    raise PasswordProtectedError(
                        "Password-protected PDFs are not supported. "
                        "Please remove the password and re-upload."
                    )

            # Sanity check: at least one page
            if len(reader.pages) == 0:
                raise EmptyDocumentError(
                    "The uploaded PDF has zero pages."
                )

        except (PasswordProtectedError, EmptyDocumentError):
            raise
        except PdfReadError as e:
            logger.warning(f"PDF preflight failed (PdfReadError): {e}")
            raise CorruptedFileError(
                "The uploaded file appears to be corrupted. "
                "Please check the file and try again."
            )
        except Exception as e:
            # Non-fatal: let the extraction stages handle it
            logger.warning(f"PDF preflight warning: {e}")

    @staticmethod
    def _pdf_stage_pypdf2(path: Path) -> Optional[str]:
        """Stage 1: Extract text using PyPDF2."""
        try:
            from PyPDF2 import PdfReader

            reader = PdfReader(str(path))
            pages_text = []

            for page_num, page in enumerate(reader.pages, start=1):
                try:
                    text = page.extract_text()
                    if text and text.strip():
                        pages_text.append(text.strip())
                except Exception as page_err:
                    logger.warning(
                        f"PyPDF2: page {page_num} of '{path.name}' failed: {page_err}"
                    )

            return "\n\n".join(pages_text) if pages_text else ""

        except PasswordProtectedError:
            raise
        except Exception as e:
            logger.warning(f"PyPDF2 extraction failed for '{path.name}': {e}")
            return ""

    @staticmethod
    def _pdf_stage_pdfplumber(path: Path) -> Optional[str]:
        """Stage 2: Extract text using pdfplumber (better layout handling)."""
        try:
            import pdfplumber
        except ImportError:
            logger.warning("pdfplumber not installed. Skipping stage 2.")
            return ""

        try:
            pages_text = []
            with pdfplumber.open(str(path)) as pdf:
                for page_num, page in enumerate(pdf.pages, start=1):
                    try:
                        text = page.extract_text()
                        if text and text.strip():
                            pages_text.append(text.strip())

                        # Also try extracting tables as text
                        tables = page.extract_tables()
                        if tables:
                            for table in tables:
                                table_text = "\n".join(
                                    "\t".join(str(cell or "") for cell in row)
                                    for row in table
                                    if row
                                )
                                if table_text.strip():
                                    pages_text.append(table_text.strip())

                    except Exception as page_err:
                        logger.warning(
                            f"pdfplumber: page {page_num} of '{path.name}' failed: {page_err}"
                        )

            return "\n\n".join(pages_text) if pages_text else ""

        except Exception as e:
            logger.warning(f"pdfplumber extraction failed for '{path.name}': {e}")
            return ""

    @staticmethod
    def _pdf_stage_ocr(path: Path) -> Optional[str]:
        """
        Stage 3: OCR extraction using pdf2image + pytesseract.

        Converts each PDF page to an image, then runs Tesseract OCR.
        Gracefully degrades if dependencies are missing.
        """
        try:
            from pdf2image import convert_from_path
            import pytesseract
        except ImportError as ie:
            logger.warning(
                f"OCR dependencies not fully installed ({ie}). "
                "Install pdf2image, pytesseract, and Pillow for OCR support. "
                "Also ensure Tesseract is installed on the system."
            )
            return ""

        try:
            logger.info(f"Running OCR on '{path.name}'...")

            # Convert PDF pages to images (200 DPI balances quality vs speed)
            try:
                images = convert_from_path(
                    str(path),
                    dpi=200,
                    fmt="png",
                    thread_count=2,
                )
            except Exception as convert_err:
                logger.warning(
                    f"pdf2image conversion failed for '{path.name}': {convert_err}. "
                    "Ensure poppler-utils is installed on the system."
                )
                return ""

            pages_text = []
            for page_num, image in enumerate(images, start=1):
                try:
                    text = pytesseract.image_to_string(image, lang="eng")
                    if text and text.strip():
                        pages_text.append(text.strip())
                    else:
                        logger.debug(f"OCR: page {page_num} yielded no text.")
                except Exception as ocr_err:
                    logger.warning(
                        f"OCR: page {page_num} of '{path.name}' failed: {ocr_err}"
                    )

            result = "\n\n".join(pages_text) if pages_text else ""
            logger.info(
                f"OCR completed for '{path.name}': {len(result)} chars "
                f"from {len(images)} pages."
            )
            return result

        except Exception as e:
            logger.warning(f"OCR pipeline failed for '{path.name}': {e}")
            return ""

    # ── DOCX Extraction ─────────────────────────────────────────────────────

    @staticmethod
    def _load_docx(path: Path) -> str:
        """
        Extract text from a DOCX file using python-docx.

        Extracts:
            - Paragraph text
            - Table cell text
        """
        logger.info(f"Extracting text from DOCX: {path.name}")
        try:
            from docx import Document as DocxDocument
        except ImportError:
            raise FileLoadError(
                "python-docx is not installed. Cannot process DOCX files."
            )

        try:
            doc = DocxDocument(str(path))
        except Exception as e:
            error_str = str(e).lower()
            if "password" in error_str or "encrypt" in error_str:
                raise PasswordProtectedError(
                    "Password-protected DOCX files are not supported. "
                    "Please remove the password and re-upload."
                )
            logger.error(f"Failed to open DOCX '{path.name}': {e}")
            raise CorruptedFileError(
                "The uploaded file appears to be corrupted or is not a valid DOCX document."
            )

        try:
            text_parts = []

            # Extract paragraphs
            for para in doc.paragraphs:
                stripped = para.text.strip()
                if stripped:
                    text_parts.append(stripped)

            # Extract table content
            for table in doc.tables:
                for row in table.rows:
                    row_text = "\t".join(
                        cell.text.strip() for cell in row.cells if cell.text.strip()
                    )
                    if row_text.strip():
                        text_parts.append(row_text.strip())

            full_text = "\n\n".join(text_parts)

            if not full_text.strip():
                raise EmptyDocumentError(
                    f"The DOCX file '{path.name}' contains no readable text content."
                )

            logger.info(
                f"DOCX '{path.name}': extracted {len(full_text)} chars "
                f"from {len(doc.paragraphs)} paragraphs + {len(doc.tables)} tables."
            )
            return full_text

        except (EmptyDocumentError, CorruptedFileError):
            raise
        except Exception as e:
            logger.error(f"Error reading DOCX content '{path.name}': {e}")
            raise CorruptedFileError(
                "The uploaded file appears to be corrupted or unreadable."
            )

    # ── TXT Extraction ──────────────────────────────────────────────────────

    @staticmethod
    def _load_txt(path: Path) -> str:
        """
        Read text from a plain text file.

        Encoding cascade: UTF-8 -> UTF-8-SIG (BOM) -> Latin-1.
        """
        logger.info(f"Reading TXT file: {path.name}")

        encodings = ["utf-8", "utf-8-sig", "latin-1", "cp1252"]

        for encoding in encodings:
            try:
                text = path.read_text(encoding=encoding)
                if text.strip():
                    logger.info(
                        f"TXT '{path.name}': read {len(text)} chars "
                        f"(encoding: {encoding})."
                    )
                    return text
            except UnicodeDecodeError:
                logger.debug(f"TXT '{path.name}': {encoding} decode failed, trying next...")
                continue
            except Exception as e:
                logger.error(f"Error reading TXT '{path.name}' with {encoding}: {e}")
                raise CorruptedFileError(
                    "The uploaded text file appears to be corrupted or unreadable."
                )

        # If all encodings succeeded but text was empty
        try:
            raw = path.read_bytes()
            if len(raw) == 0:
                raise EmptyDocumentError(
                    f"The TXT file '{path.name}' is empty."
                )
        except EmptyDocumentError:
            raise

        raise EmptyDocumentError(
            f"The TXT file '{path.name}' contains no readable text content."
        )